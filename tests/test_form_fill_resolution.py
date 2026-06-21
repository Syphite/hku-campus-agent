"""Tests for approved-data key resolution and DOCX row detection."""

import os
import tempfile
import unittest
from unittest.mock import MagicMock

from agent.application.docx_filler import (
    _find_header_row,
    _fill_repeating_list,
    _header_column_map,
    _is_row_label_cell,
    _next_empty_data_row,
    _row_cells,
    fill_docx_form,
)
from agent.application.form_ai import resolve_approved_repeating_lists

FIXTURE_DOCX = os.path.join(
    os.path.dirname(__file__),
    "fixtures",
    "DHCFS_Application_Form_2627.docx",
)


class ResolveApprovedRepeatingListsTests(unittest.TestCase):
    def test_maps_mismatched_pending_key_to_schema_key(self):
        table_schema = {
            "repeating_lists": [
                {
                    "key": "awards_achievements",
                    "label": "Awards / Achievements",
                    "list_kind": "awards",
                    "table_index": 2,
                    "column_headers": ["Date", "Award", "Organization"],
                }
            ]
        }
        pending = {
            "awards": [
                {"dates": "06/24", "role": "Dean's List", "organization": "HKU"},
            ]
        }
        gap_queue = [
            {
                "key": "awards",
                "label": "Awards / Achievements",
                "schema": {
                    "key": "awards",
                    "list_kind": "awards",
                    "table_index": 2,
                },
            }
        ]

        resolved, _warnings = resolve_approved_repeating_lists(table_schema, pending, gap_queue)

        self.assertIn("awards_achievements", resolved)
        self.assertEqual(len(resolved["awards_achievements"]), 1)
        self.assertEqual(resolved["awards_achievements"][0]["role"], "Dean's List")


class DocxRowDetectionTests(unittest.TestCase):
    def test_row_label_cells_recognized(self):
        self.assertTrue(_is_row_label_cell("C1"))
        self.assertTrue(_is_row_label_cell("f2"))
        self.assertFalse(_is_row_label_cell("Award"))

    def test_next_empty_data_row_skips_row_label_column(self):
        table = MagicMock()
        table.rows = [MagicMock(), MagicMock(), MagicMock()]

        header_cells = [MagicMock(), MagicMock(), MagicMock()]
        header_cells[0].text = "Date"
        header_cells[1].text = "Award"
        header_cells[2].text = "Organization"

        data_cells = [MagicMock(), MagicMock(), MagicMock()]
        data_cells[0].text = "C1"
        data_cells[1].text = ""
        data_cells[2].text = ""

        table.rows[0].cells = header_cells
        table.rows[1].cells = data_cells

        def row_cells_side_effect(_table, row_index):
            if row_index == 0:
                return header_cells
            if row_index == 1:
                return data_cells
            return []

        column_map = {"dates": 1, "role": 2}

        import agent.application.docx_filler as docx_filler

        original = docx_filler._row_cells
        docx_filler._row_cells = row_cells_side_effect
        try:
            match = _next_empty_data_row(table, 0, column_map, 5)
        finally:
            docx_filler._row_cells = original

        self.assertIsNotNone(match)
        row_index, _cells = match
        self.assertEqual(row_index, 1)

    def test_find_header_row_requires_two_header_matches(self):
        table = MagicMock()
        table.rows = [MagicMock(), MagicMock()]

        instruction_cells = [MagicMock()]
        instruction_cells[0].text = "Section C: list awards below with Date and Award columns"

        header_cells = [MagicMock(), MagicMock(), MagicMock()]
        header_cells[0].text = "Date (MM/YY)"
        header_cells[1].text = "Name of Award"
        header_cells[2].text = "Organization"

        table.rows[0].cells = instruction_cells
        table.rows[1].cells = header_cells

        import agent.application.docx_filler as docx_filler

        original = docx_filler._row_cells

        def row_cells_side_effect(_table, row_index):
            if row_index == 0:
                return instruction_cells
            if row_index == 1:
                return header_cells
            return []

        docx_filler._row_cells = row_cells_side_effect
        try:
            match = _find_header_row(
                table,
                ["Date (MM/YY)", "Name of Award", "Organization"],
            )
        finally:
            docx_filler._row_cells = original

        self.assertIsNotNone(match)
        row_index, _cells = match
        self.assertEqual(row_index, 1)


class DhcfsAwardsTableTests(unittest.TestCase):
    def test_awards_header_maps_past_row_label_column(self):
        if not os.path.exists(FIXTURE_DOCX):
            self.skipTest("DHCFS fixture DOCX not available")

        from docx import Document

        table = Document(FIXTURE_DOCX).tables[3]
        headers = ["Year (MM/YY)", "Name of the Award", "Awarding Unit"]
        item_fields = {
            "dates": "Year (MM/YY)",
            "role": "Name of the Award",
            "organization": "Awarding Unit",
        }
        header_match = _find_header_row(table, headers)
        self.assertIsNotNone(header_match)
        _row_index, header_cells = header_match

        column_map = _header_column_map(header_cells, headers, item_fields)
        self.assertEqual(column_map, {"dates": 1, "role": 2, "organization": 3})

    def test_awards_rows_fill_in_fixture_docx(self):
        if not os.path.exists(FIXTURE_DOCX):
            self.skipTest("DHCFS fixture DOCX not available")

        schema = {
            "repeating_lists": [
                {
                    "key": "awards",
                    "table_index": 3,
                    "column_headers": ["Year (MM/YY)", "Name of the Award", "Awarding Unit"],
                    "item_fields": {
                        "dates": "Year (MM/YY)",
                        "role": "Name of the Award",
                        "organization": "Awarding Unit",
                    },
                    "max_rows": 5,
                }
            ]
        }
        filled_data = {
            "repeating_lists": {
                "awards": [
                    {"dates": "06/24", "role": "Dean's List", "organization": "HKU"},
                    {"dates": "05/23", "role": "Hackathon Winner", "organization": "HKSTP"},
                ]
            }
        }

        with tempfile.TemporaryDirectory() as tmp:
            output_path = os.path.join(tmp, "filled.docx")
            _, fill_report = fill_docx_form(FIXTURE_DOCX, filled_data, schema, output_path)
            self.assertEqual(fill_report["repeating_lists"]["awards"], 2)

            from docx import Document

            table = Document(output_path).tables[3]
            row1 = [cell.text.strip() for cell in _row_cells(table, 1)]
            row2 = [cell.text.strip() for cell in _row_cells(table, 2)]
            self.assertEqual(row1[1:], ["06/24", "Dean's List", "HKU"])
            self.assertEqual(row2[1:], ["05/23", "Hackathon Winner", "HKSTP"])


if __name__ == "__main__":
    unittest.main()
