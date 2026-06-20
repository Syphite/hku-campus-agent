"""Extract structured table JSON from DOCX application forms."""

from __future__ import annotations

import json
import logging

from agent.application.cell_utils import dedupe_row_cells

logger = logging.getLogger(__name__)


def _extract_content_controls(doc) -> list[dict]:
    """Scan document body for structured document tags (content controls)."""
    controls = []
    try:
        from docx.oxml.ns import qn

        body = doc.element.body
        for index, sdt in enumerate(body.iter(qn("w:sdt"))):
            texts = [node.text for node in sdt.iter(qn("w:t")) if node.text]
            label = " ".join(texts).replace("\n", " ").strip()
            tag_node = sdt.find(qn("w:sdtPr"))
            tag_value = ""
            if tag_node is not None:
                tag_el = tag_node.find(qn("w:tag"))
                if tag_el is not None:
                    tag_value = tag_el.get(qn("w:val"), "")
            if label or tag_value:
                controls.append({
                    "index": index,
                    "label": label[:200],
                    "tag": tag_value,
                })
    except Exception as exc:
        logger.debug("Content control scan skipped: %s", exc)
    return controls


def extract_form_schema(file_path: str) -> str:
    """
    Parse a DOCX form into structured JSON for LLM schema analysis.

    Returns a JSON string with tables, paragraphs, and content controls.
    """
    from docx import Document

    doc = Document(file_path)
    tables = []

    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            raw_cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            cells = dedupe_row_cells(raw_cells)
            rows.append({
                "cells": [cell if cell else "[Empty]" for cell in cells]
            })
        tables.append({"table_index": table_index, "rows": rows})

    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.replace("\n", " ").strip()
        if text:
            paragraphs.append({"paragraph_index": index, "text": text})

    content_controls = _extract_content_controls(doc)

    payload = {
        "format": "docx",
        "source_path": file_path,
        "tables": tables,
        "paragraphs": paragraphs,
        "content_controls": content_controls,
        "acroform_widgets": [],
    }
    logger.info(
        "DOCX schema extracted: %s tables, %s paragraphs, %s content controls",
        len(tables),
        len(paragraphs),
        len(content_controls),
    )
    return json.dumps(payload, ensure_ascii=False, indent=2)
