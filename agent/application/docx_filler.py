"""Anchor-based DOCX form filling."""

from __future__ import annotations

import logging
import os
import re
import shutil

from agent.application.cell_utils import is_empty_cell, normalize_text, texts_match
from agent.application.fill_geometry import resolve_description_target, resolve_docx_table_target

logger = logging.getLogger(__name__)

_ROW_LABEL_RE = re.compile(r"^[A-Z]\d+$", re.IGNORECASE)


def _is_row_label_cell(text: str) -> bool:
    compact = normalize_text(text).replace(" ", "")
    return bool(_ROW_LABEL_RE.match(compact))


def _set_cell_text(cell, value: str) -> None:
    text = str(value or "").strip()
    if not text:
        return
    cell.text = text


def _row_cells(table, row_index: int):
    row = table.rows[row_index]
    seen = set()
    cells = []
    for cell in row.cells:
        cell_id = id(cell._tc)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        cells.append(cell)
    return cells


def _find_anchor(table, anchor_label: str):
    target = normalize_text(anchor_label)
    for row_index, row in enumerate(table.rows):
        cells = _row_cells(table, row_index)
        for col_index, cell in enumerate(cells):
            if texts_match(anchor_label, cell.text):
                return row_index, col_index, cell, cells
    return None


def _fill_at_location(table, row_index: int, col_index: int, cells, fill_location: str, value: str) -> bool:
    """Fill a table cell using layout heuristics (right if empty, else below)."""
    if not value:
        return False

    fill_location = (fill_location or "auto").lower()

    if fill_location == "right" and col_index + 1 < len(cells):
        if is_empty_cell(cells[col_index + 1].text):
            _set_cell_text(cells[col_index + 1], value)
            return True

    if fill_location == "below" and row_index + 1 < len(table.rows):
        below_cells = _row_cells(table, row_index + 1)
        if col_index < len(below_cells) and is_empty_cell(below_cells[col_index].text):
            _set_cell_text(below_cells[col_index], value)
            return True
        for cell in below_cells:
            if is_empty_cell(cell.text):
                _set_cell_text(cell, value)
                return True

    target = resolve_docx_table_target(table, row_index, col_index, cells)
    if target is not None:
        _set_cell_text(target, value)
        return True
    return False


def _find_header_row(table, headers: list[str]):
    normalized_headers = [normalize_text(header) for header in headers if header]
    if len(normalized_headers) < 2:
        return None

    best_match = None
    best_score = 0
    for row_index, row in enumerate(table.rows):
        cells = _row_cells(table, row_index)
        row_texts = [normalize_text(cell.text) for cell in cells]
        if len([text for text in row_texts if text]) <= 1:
            continue
        score = 0
        for header in normalized_headers:
            if any(texts_match(header, text) for text in row_texts):
                score += 1
        if score >= 2 and score > best_score:
            best_match = (row_index, cells)
            best_score = score
    return best_match


def _header_column_map(cells, headers: list[str], item_fields: dict) -> dict:
    mapping = {}
    for logical_key, header_label in (item_fields or {}).items():
        for col_index, cell in enumerate(cells):
            if _is_row_label_cell(cell.text):
                continue
            if not normalize_text(cell.text):
                continue
            if texts_match(header_label, cell.text):
                mapping[logical_key] = col_index
                break
    return mapping


def _is_description_row(cells) -> bool:
    row_text = " ".join(normalize_text(cell.text) for cell in cells)
    return "description" in row_text or "in 50 words" in row_text


def _next_empty_data_row(table, start_row: int, column_map: dict, max_rows: int, *, description_row: bool = False):
    data_columns = sorted(set(column_map.values()))
    if not data_columns:
        return None

    rows_checked = 0
    for row_index in range(start_row + 1, len(table.rows)):
        if rows_checked >= max_rows:
            break
        cells = _row_cells(table, row_index)
        if description_row and _is_description_row(cells):
            continue
        values = []
        for col_index in data_columns:
            if col_index < len(cells):
                values.append(cells[col_index].text)
        if not values:
            continue
        if all(is_empty_cell(value) for value in values):
            rows_checked += 1
            return row_index, cells
    return None


def _description_target_cell(table, data_row_index: int, list_schema: dict | None = None):
    """Find the cell for a list-item description using table layout heuristics."""
    del list_schema  # table geometry is resolved from structure, not form-specific rules
    return resolve_description_target(table, data_row_index, is_description_row=_is_description_row)


def _fill_repeating_list(table, list_schema: dict, items: list[dict]) -> int:
    headers = list_schema.get("column_headers") or []
    item_fields = list_schema.get("item_fields") or {}
    max_rows = int(list_schema.get("max_rows") or len(items) or 5)
    header_match = _find_header_row(table, headers)
    if not header_match:
        logger.warning("Could not find header row for repeating list %s", list_schema.get("key"))
        return 0

    header_row_index, header_cells = header_match
    column_map = _header_column_map(header_cells, headers, item_fields)
    if not column_map:
        logger.warning("Could not map columns for repeating list %s", list_schema.get("key"))
        return 0

    use_description_row = bool(list_schema.get("description_row"))
    filled = 0
    last_row = header_row_index
    for item in items[:max_rows]:
        target = _next_empty_data_row(
            table,
            last_row,
            column_map,
            max_rows,
            description_row=use_description_row,
        )
        if not target:
            break
        row_index, cells = target
        for logical_key, col_index in column_map.items():
            if use_description_row and logical_key == "description":
                continue
            if col_index < len(cells):
                value = item.get(logical_key, "")
                _set_cell_text(cells[col_index], value)
        if use_description_row:
            description = item.get("description") or ""
            if description:
                target_cell = _description_target_cell(table, row_index, list_schema)
                if target_cell is not None:
                    _set_cell_text(target_cell, description)
        filled += 1
        last_row = row_index
    if items and filled == 0:
        logger.warning(
            "Repeating list %s had %d items but filled 0 rows (table_index=%s, headers=%s, column_map=%s)",
            list_schema.get("key"),
            len(items),
            list_schema.get("table_index"),
            headers,
            column_map,
        )
    return filled


def fill_docx_form(original_path: str, filled_data: dict, schema: dict, output_path: str | None = None) -> tuple[str, dict]:
    from docx import Document

    if not os.path.exists(original_path):
        raise FileNotFoundError(original_path)

    target_path = output_path or original_path.replace(".docx", "_filled.docx")
    if target_path == original_path:
        target_path = original_path.replace(".docx", "_filled.docx")
    shutil.copy2(original_path, target_path)

    doc = Document(target_path)
    tables = doc.tables
    fill_report = {"repeating_lists": {}}

    for field in schema.get("simple_fields", []):
        table_index = int(field.get("table_index", 0))
        if table_index >= len(tables):
            continue
        anchor = _find_anchor(tables[table_index], field.get("anchor_label", ""))
        if not anchor:
            continue
        row_index, col_index, _, cells = anchor
        value = (filled_data.get("simple_fields") or {}).get(field.get("key"), "")
        _fill_at_location(
            tables[table_index],
            row_index,
            col_index,
            cells,
            field.get("fill_location", "auto"),
            value,
        )

    for field in schema.get("booleans", []):
        table_index = int(field.get("table_index", 0))
        if table_index >= len(tables):
            continue
        anchor = _find_anchor(tables[table_index], field.get("anchor_label", ""))
        if not anchor:
            continue
        row_index, col_index, _, cells = anchor
        raw_value = (filled_data.get("booleans") or {}).get(field.get("key"))
        if raw_value is None:
            continue
        value = "Yes" if raw_value else "No"
        _fill_at_location(
            tables[table_index],
            row_index,
            col_index,
            cells,
            field.get("fill_location", "auto"),
            value,
        )

    for field in schema.get("long_text", []):
        table_index = int(field.get("table_index", 0))
        value = (filled_data.get("long_text") or {}).get(field.get("key"), "")
        if not value:
            continue
        if table_index < len(tables):
            anchor = _find_anchor(tables[table_index], field.get("anchor_label", ""))
            if anchor:
                row_index, col_index, _, cells = anchor
                _fill_at_location(
                    tables[table_index],
                    row_index,
                    col_index,
                    cells,
                    field.get("fill_location", "auto"),
                    value,
                )
                continue
        for paragraph in doc.paragraphs:
            if texts_match(field.get("anchor_label", ""), paragraph.text):
                paragraph.text = f"{paragraph.text}\n{value}".strip()
                break

    for field in schema.get("repeating_lists", []):
        table_index = int(field.get("table_index", 0))
        if table_index >= len(tables):
            continue
        items = (filled_data.get("repeating_lists") or {}).get(field.get("key"), [])
        if items:
            field_key = field.get("key")
            filled_count = _fill_repeating_list(tables[table_index], field, items)
            fill_report["repeating_lists"][field_key] = filled_count

    doc.save(target_path)
    logger.info("Filled DOCX saved to %s", target_path)
    return target_path, fill_report


def fill_free_fields_docx(
    docx_path: str,
    free_field_defs: list[dict],
    free_field_values: dict,
) -> int:
    """Fill paragraph/content-control free fields in an already-saved DOCX."""
    from docx import Document

    if not free_field_values:
        return 0

    doc = Document(docx_path)
    filled = 0

    for field in free_field_defs or []:
        key = field.get("key")
        value = str((free_field_values or {}).get(key) or "").strip()
        if not value:
            continue

        paragraph_index = field.get("paragraph_index")
        if paragraph_index is not None:
            try:
                idx = int(paragraph_index)
                if 0 <= idx < len(doc.paragraphs):
                    paragraph = doc.paragraphs[idx]
                    if texts_match(field.get("anchor_label", ""), paragraph.text):
                        paragraph.text = f"{paragraph.text}\n{value}".strip()
                    else:
                        paragraph.text = value
                    filled += 1
                    continue
            except (TypeError, ValueError):
                pass

        anchor = field.get("anchor_label") or field.get("question_text", "")
        for paragraph in doc.paragraphs:
            if texts_match(anchor, paragraph.text):
                paragraph.text = f"{paragraph.text}\n{value}".strip()
                filled += 1
                break

    doc.save(docx_path)
    return filled
