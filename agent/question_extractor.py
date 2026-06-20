"""
agent/question_extractor.py
Extracts essay/application questions from uploaded PDF or DOCX forms.
"""

import io
import json
import logging
import os
from typing import Optional

from dotenv import load_dotenv
from openai import AzureOpenAI

load_dotenv()
logger = logging.getLogger(__name__)


def _get_openai_client() -> Optional[AzureOpenAI]:
    endpoint = os.environ.get("AZURE_OPENAI_ENDPOINT")
    api_key = os.environ.get("AZURE_OPENAI_API_KEY")
    if not endpoint or not api_key:
        logger.warning("Azure OpenAI is not configured; cannot extract form questions")
        return None

    return AzureOpenAI(
        azure_endpoint=endpoint,
        api_key=api_key,
        api_version="2024-12-01-preview"
    )


def _extract_docx_text(source) -> str:
    from docx import Document

    doc = Document(source)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    return "\n".join(parts).strip()


def _extract_pdf_text(source) -> str:
    from pypdf import PdfReader

    if isinstance(source, (bytes, bytearray)):
        reader = PdfReader(io.BytesIO(source))
    else:
        reader = PdfReader(source)
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def extract_text_from_application_file(file_path: str) -> str:
    """Extract plain text from a saved PDF or DOCX application form."""
    path_lower = (file_path or "").lower()
    try:
        if path_lower.endswith(".pdf"):
            return _extract_pdf_text(file_path)
        if path_lower.endswith(".docx"):
            return _extract_docx_text(file_path)
    except Exception as e:
        logger.error(f"Application form text extraction failed for {file_path}: {e}")
    return ""


def _extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    filename_lower = (filename or "").lower()

    if filename_lower.endswith(".pdf"):
        try:
            return _extract_pdf_text(file_bytes)
        except Exception as e:
            logger.error(f"PDF question extraction failed for {filename}: {e}")
            return ""

    if filename_lower.endswith(".docx"):
        try:
            return _extract_docx_text(io.BytesIO(file_bytes))
        except Exception as e:
            logger.error(f"DOCX question extraction failed for {filename}: {e}")
            return ""

    logger.warning(f"Unsupported form upload type: {filename}")
    return ""


def _normalize_questions(values: list) -> list[dict]:
    questions = []
    seen = set()
    for idx, value in enumerate(values, start=1):
        if isinstance(value, dict):
            text = str(value.get("text") or value.get("question") or "").strip()
        else:
            text = str(value).strip()

        key = text.lower()
        if text and key not in seen:
            seen.add(key)
            questions.append({"id": f"q{len(questions) + 1}", "text": text})

    return questions[:10]


def _fallback_questions(raw_text: str) -> list[dict]:
    values = []
    for line in (raw_text or "").splitlines():
        text = line.strip().lstrip("-*0123456789. )\t")
        if not text:
            continue
        if "?" in text or len(text.split()) >= 6:
            values.append(text)
    return _normalize_questions(values)


def extract_questions_from_file(file_bytes: bytes, filename: str) -> list[dict]:
    """
    Extract essay/application questions from a PDF or DOCX form.
    Returns [{"id": "q1", "text": "..."}].
    """
    raw_text = _extract_text_from_file(file_bytes, filename)
    if not raw_text:
        return []

    client = _get_openai_client()
    if not client:
        return _fallback_questions(raw_text)

    prompt = f"""
You are parsing a scholarship or university application form.
Extract only essay questions or prompts the student must answer.

Ignore basic fields such as name, email, phone, student ID, faculty, programme,
GPA, nationality, signature, date, checkbox labels, upload instructions, and
document headings.

Return JSON only in this exact shape:
{{"questions": [{{"id": "q1", "text": "Question text here"}}]}}

Extract at most 10 questions.

Form text:
{raw_text[:12000]}
"""

    try:
        response = client.chat.completions.create(
            model=os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4o"),
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            max_tokens=1500,
            temperature=0.1
        )
        parsed = json.loads(response.choices[0].message.content)
        values = parsed.get("questions", []) if isinstance(parsed, dict) else []
        return _normalize_questions(values)
    except json.JSONDecodeError as e:
        logger.error(f"Question extraction JSON parse error: {e}")
        return _fallback_questions(raw_text)
    except Exception as e:
        logger.error(f"Question extraction failed: {e}")
        return _fallback_questions(raw_text)
