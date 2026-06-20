"""
agent/form_filler.py
Fills out scholarship application forms (PDF or DOCX) using the student's profile.
Automatically detects file type and applies the appropriate filling strategy.
"""
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# =============================================================================
# PDF Filling Logic (AcroForm)
# =============================================================================
def _fill_pdf(input_path: str, output_path: str, profile: dict, scholarship: dict) -> bool:
    """Fills AcroForm text fields in a PDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.error("PyMuPDF (fitz) is not installed. Run: pip install PyMuPDF")
        return False

    try:
        doc = fitz.open(input_path)
        
        # Build a flat mapping of expected form field names to values
        academic = profile.get("academic", {})
        financial = profile.get("financial", {})
        
        field_mapping = {
            "name": profile.get("name", ""),
            "full_name": profile.get("name", ""),
            "student_name": profile.get("name", ""),
            "email": profile.get("email", ""),
            "phone": profile.get("phone", ""),
            "student_id": profile.get("student_id", ""),
            "faculty": academic.get("faculty", ""),
            "programme": academic.get("programme", ""),
            "major": academic.get("programme", ""),
            "year_of_study": str(academic.get("year_of_study", "")),
            "year": str(academic.get("year_of_study", "")),
            "gpa": str(academic.get("gpa", "")),
            "cgpa": str(academic.get("gpa", "")),
            "nationality": academic.get("nationality", {}).get("country_of_origin", ""),
            "scholarship_name": scholarship.get("name", ""),
        }

        filled_count = 0
        for widget in doc.widgets():
            if not widget.field_name:
                continue
                
            # Normalize field name for matching
            field_name = widget.field_name.lower().replace(" ", "_").replace("-", "_")
            
            # Check if this field matches something in our profile
            for key, value in field_mapping.items():
                if key in field_name and value:
                    widget.field_value = str(value)
                    widget.update()  # Apply the change visually
                    filled_count += 1
                    break

        doc.save(output_path)
        doc.close()
        
        logger.info(f"Successfully filled {filled_count} fields in PDF: {output_path}")
        return True

    except Exception as e:
        logger.error(f"PDF filling failed for {input_path}: {e}")
        return False


# =============================================================================
# DOCX Filling Logic (Placeholder Replacement)
# =============================================================================
def _fill_docx(input_path: str, output_path: str, profile: dict, scholarship: dict) -> bool:
    """Fills DOCX templates by replacing {{placeholder}} tags in paragraphs and tables."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx is not installed. Run: pip install python-docx")
        return False

    try:
        doc = Document(input_path)
        
        # Build replacement dictionary with {{placeholder}} syntax
        academic = profile.get("academic", {})
        financial = profile.get("financial", {})
        
        replacements = {
            "{{name}}": profile.get("name", ""),
            "{{full_name}}": profile.get("name", ""),
            "{{email}}": profile.get("email", ""),
            "{{phone}}": profile.get("phone", ""),
            "{{student_id}}": profile.get("student_id", ""),
            "{{faculty}}": academic.get("faculty", ""),
            "{{programme}}": academic.get("programme", ""),
            "{{major}}": academic.get("programme", ""),
            "{{year_of_study}}": str(academic.get("year_of_study", "")),
            "{{gpa}}": str(academic.get("gpa", "")),
            "{{cgpa}}": str(academic.get("gpa", "")),
            "{{nationality}}": academic.get("nationality", {}).get("country_of_origin", ""),
            "{{scholarship_name}}": scholarship.get("name", ""),
            "{{cover_letter}}": scholarship.get("drafted_cover_letter", ""), # If drafter passed it in
        }

        replaced_count = 0

        # 1. Replace in paragraphs
        for para in doc.paragraphs:
            for key, val in replacements.items():
                if key in para.text and val:
                    para.text = para.text.replace(key, str(val))
                    replaced_count += 1

        # 2. Replace in tables (crucial for many application forms)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, val in replacements.items():
                            if key in para.text and val:
                                para.text = para.text.replace(key, str(val))
                                replaced_count += 1

        doc.save(output_path)
        logger.info(f"Successfully replaced {replaced_count} placeholders in DOCX: {output_path}")
        return True

    except Exception as e:
        logger.error(f"DOCX filling failed for {input_path}: {e}")
        return False


# =============================================================================
# Main Entry Point
# =============================================================================
def fill_application_form(input_path: str, output_path: str, profile: dict, scholarship: dict) -> bool:
    """
    Detects file type and routes to the appropriate filler.
    
    Args:
        input_path: Path to the blank application form (PDF or DOCX)
        output_path: Path where the filled form will be saved
        profile: Student profile dict
        scholarship: Scholarship details dict
        
    Returns:
        bool: True if successful, False otherwise
    """
    if not os.path.exists(input_path):
        logger.error(f"Input form not found: {input_path}")
        return False

    ext = os.path.splitext(input_path)[1].lower()
    
    if ext == ".pdf":
        return _fill_pdf(input_path, output_path, profile, scholarship)
    elif ext in [".docx", ".doc"]:
        return _fill_docx(input_path, output_path, profile, scholarship)
    else:
        logger.error(f"Unsupported file format for form filling: {ext}")
        return False
